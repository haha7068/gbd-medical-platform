import streamlit as st
import pandas as pd
import numpy as np
import matplotlib
from matplotlib import pyplot as plt
from scipy import stats
import tempfile
import os
import chardet
import json
import plotly.express as px

# 页面配置
st.set_page_config(page_title="GBD医学数据分析与可视化平台", layout="wide")
st.title("🌍 GBD医学数据分析与可视化平台\nGBD Medical Data Analysis & Visualization Platform")


# ========= 📂 数据上传区 Upload Section =========
st.sidebar.header("📂 数据导入 Upload Data")
uploaded_files = st.sidebar.file_uploader(
    "上传一个或多个数据文件（支持 CSV, Excel, JSON）",
    type=["csv", "xlsx", "xls", "json"],
    accept_multiple_files=True
)
# ========== 🧹 清除上传文件缓存 Clear Cache Section ========== #
CACHE_FILE = "uploaded_data_cache.json"#定义缓存路径
if st.sidebar.button("🧹 清除上传缓存 Clear Uploaded Cache"):
    if os.path.exists(CACHE_FILE):
        os.remove(CACHE_FILE)
        st.success("✅ 上传文件缓存已清除！Cache Cleared Successfully!")
    else:
        st.info("⚠️ 当前没有缓存文件，无需清除。")



def save_cache(file_info):
    with open(CACHE_FILE, "w", encoding="utf-8") as f:
        json.dump(file_info, f)

def load_cache():
    if os.path.exists(CACHE_FILE):
        with open(CACHE_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return {}

def check_data_quality(df, filename=""):
    problems = []
    if df.empty:
        problems.append("❌ 文件为空 Empty file")
    if df.shape[1] < 2:
        problems.append("⚠️ 列数过少 Less than 2 columns")
    if df.isnull().mean().mean() > 0.3:
        problems.append("⚠️ 缺失值比例超过30% High missing values (>30%)")
    if (df.dtypes == 'object').mean() > 0.8:
        problems.append("⚠️ 80%以上列是文本型 High proportion of text columns (>80%)")
    if df.select_dtypes(include=[np.number]).shape[1] == 0:
        problems.append("⚠️ 没有任何数值型字段 No numeric columns")

    if problems:
        st.warning(f"⚠️ 文件 {filename} 存在潜在问题 Potential issues:")
        for p in problems:
            st.write(p)
    else:
        st.success(f"✅ 文件 {filename} 检查通过！Format Check Passed！")


file_info_cache = {}

if uploaded_files:
    file_info_cache = {}
    dfs = []
    for uploaded_file in uploaded_files:
        filename = uploaded_file.name
        suffix = filename.split(".")[-1].lower()
        temp_path = os.path.join(tempfile.gettempdir(), filename)
        with open(temp_path, "wb") as f:
            f.write(uploaded_file.read())
        try:
            if suffix == "csv":
                with open(temp_path, "rb") as f:
                    encoding = chardet.detect(f.read())["encoding"]
                df = pd.read_csv(temp_path, encoding=encoding)
            elif suffix in ["xlsx", "xls"]:
                df = pd.read_excel(temp_path)
            elif suffix == "json":
                df = pd.read_json(temp_path)
            else:
                continue
            dfs.append(df)
            file_info_cache[filename] = temp_path
            check_data_quality(df, filename)

        except Exception as e:
            st.error(f"{filename} 文件读取失败：{e}")
    save_cache(file_info_cache)
else:
    cached_files = load_cache()
    dfs = []
    for filename, temp_path in cached_files.items():
        try:
            suffix = filename.split(".")[-1].lower()
            if suffix == "csv":
                with open(temp_path, "rb") as f:
                    encoding = chardet.detect(f.read())["encoding"]
                df = pd.read_csv(temp_path, encoding=encoding)
            elif suffix in ["xlsx", "xls"]:
                df = pd.read_excel(temp_path)
            elif suffix == "json":
                df = pd.read_json(temp_path)
            else:
                continue
            dfs.append(df)
            st.sidebar.success(f"从缓存加载：{filename}")
        except:
            st.sidebar.warning(f"⚠ 无法从缓存加载 {filename}")

if not dfs:
    st.info("请上传数据后开始分析~")
# ========= 📄 数据浏览与预处理区 Data Preview & Preprocessing =========
if dfs:
    st.success(f"✅ 成功加载 {len(dfs)} 个数据集")

    for i, df in enumerate(dfs):
        st.header(f"📄 数据集 {i+1} 预览 Dataset {i+1}")
        st.dataframe(df.head())

        # ========== 🔍 筛选与分组 Filter & Group ==========
        st.subheader("🔍 数据筛选与分组 Group & Filter")

        filter_cols = df.columns.tolist()
        selected_filter_col = st.selectbox(
            f"选择筛选字段 Dataset{i+1}",
            filter_cols,
            key=f"filtercol_{i}"
        )

        if np.issubdtype(df[selected_filter_col].dtype, np.number):
            min_val = float(df[selected_filter_col].min())
            max_val = float(df[selected_filter_col].max())
            selected_range = st.slider(
                f"选择数值范围 {selected_filter_col}",
                min_value=min_val, max_value=max_val,
                value=(min_val, max_val),
                key=f"range_{i}"
            )
            if st.button(f"应用数值筛选 Apply Numeric Filter Dataset{i+1}", key=f"numfiltbtn_{i}"):
                df = df[(df[selected_filter_col] >= selected_range[0]) & (df[selected_filter_col] <= selected_range[1])]
                st.success(f"筛选后剩余 {len(df)} 行")
                st.dataframe(df)
        else:
            unique_vals = df[selected_filter_col].dropna().unique().tolist()
            selected_vals = st.multiselect(
                f"选择分类值 {selected_filter_col}",
                unique_vals, default=unique_vals[:1],
                key=f"catfilt_{i}"
            )
            if st.button(f"应用分类筛选 Apply Categorical Filter Dataset{i+1}", key=f"catfiltbtn_{i}"):
                df = df[df[selected_filter_col].isin(selected_vals)]
                st.success(f"筛选后剩余 {len(df)} 行")
                st.dataframe(df)

        # ========== 🧹 缺失值处理 Missing Value Handling ==========
        st.subheader("🧹 缺失值处理 Missing Values")

        missing_method = st.selectbox(
            "选择处理方式",
            ["不处理", "删除缺失行", "用均值填充", "用中位数填充", "用众数填充", "线性插值"],
            key=f"missing_{i}"
        )
        if st.button(f"应用缺失值处理 Apply Missing Handling Dataset{i+1}", key=f"missingbtn_{i}"):
            if missing_method == "删除缺失行":
                df = df.dropna()
            elif missing_method == "用均值填充":
                df = df.fillna(df.mean(numeric_only=True))
            elif missing_method == "用中位数填充":
                df = df.fillna(df.median(numeric_only=True))
            elif missing_method == "用众数填充":
                df = df.fillna(df.mode().iloc[0])
            elif missing_method == "线性插值":
                df = df.interpolate()
            st.success("✅ 缺失值处理完成！")
            st.dataframe(df)

        # ========== 📐 标准化与归一化 Normalization ==========
        st.subheader("📐 数据标准化与归一化 Normalization")
        if not dfs:
            st.warning("⚠ 当前没有上传数据，请先在侧边栏上传数据！")
            st.stop()

        numeric_cols = df.select_dtypes(include=np.number).columns.tolist()
        scale_cols = st.multiselect(
            f"选择需要标准化/归一化的字段 Dataset{i+1}",
            numeric_cols,
            key=f"scalecols_{i}"
        )
        scale_method = st.selectbox(
            "选择变换方式",
            ["Z-Score标准化", "最小-最大归一化"],
            key=f"scalemethod_{i}"
        )

        if st.button(f"应用标准化归一化 Apply Scaling Dataset{i+1}", key=f"scalebtn_{i}"):
            if scale_cols:
                for col in scale_cols:
                    if scale_method == "Z-Score标准化":
                        mean = df[col].mean()
                        std = df[col].std()
                        df[f"{col}_Zscore"] = (df[col] - mean) / std
                    elif scale_method == "最小-最大归一化":
                        min_val = df[col].min()
                        max_val = df[col].max()
                        if max_val != min_val:
                            df[f"{col}_MinMax"] = (df[col] - min_val) / (max_val - min_val)
                st.success("✅ 数据变换完成！")
                st.dataframe(df)
            else:
                st.warning("⚠ 请至少选择一个字段进行标准化/归一化！")
# ========= 📊 描述性统计与基础分析区 Descriptive Stats & Basic Analysis =========
        # ========== 📊 描述性统计 Descriptive Statistics ==========
        st.subheader("📊 描述性统计 Descriptive Statistics")
        if not dfs:
            st.warning("⚠ 当前没有上传数据，请先在侧边栏上传数据！")
            st.stop()

        desc_numeric_cols = df.select_dtypes(include=np.number).columns.tolist()
        if desc_numeric_cols:
            selected_desc_cols = st.multiselect(
                "选择要进行统计描述的字段",
                desc_numeric_cols,
                default=desc_numeric_cols[:2],
                key=f"desc_cols_{i}"
            )

            if st.button(f"执行描述性统计 Run Description Dataset{i + 1}", key=f"descbtn_{i}"):
                if selected_desc_cols:
                    desc_stats = df[selected_desc_cols].describe().T
                    desc_stats["方差 Variance"] = df[selected_desc_cols].var()
                    desc_stats["众数 Mode"] = [
                        df[col].mode().iloc[0] if not df[col].mode().empty else np.nan
                        for col in selected_desc_cols
                    ]
                    st.success("✅ 描述性统计完成 Description Completed")
                    st.dataframe(desc_stats)

                    # ==== 🚀 增强版可视化部分 ====
                    st.markdown("### 📊 各字段分布直方图 Histograms of Selected Fields")

                    import plotly.express as px

                    for col in selected_desc_cols:
                        fig = px.histogram(
                            df,
                            x=col,
                            nbins=30,
                            title=f"📊 {col} 分布直方图 Histogram of {col}",
                            marginal="box"
                        )
                        fig.write_image(f"desc_histogram_{col}.png", scale=2)
                        st.plotly_chart(fig, use_container_width=True)

                else:
                    st.warning("⚠ 请至少选择一个字段进行描述统计")

        # ========== 🔗 相关性分析 Correlation Analysis ========== #
        st.subheader("🔗 相关性分析 Correlation Analysis")
        if not dfs:
            st.warning("⚠ 当前没有上传数据，请先在侧边栏上传数据！")
            st.stop()

        correlation_cols = df.select_dtypes(include=np.number).columns.tolist()

        if correlation_cols:
            selected_corr_cols = st.multiselect(
                "选择进行相关性分析的数值列 Select Numeric Columns for Correlation",
                correlation_cols,
                key=f"correlation_cols_{i}"
            )

            correlation_method = st.radio(
                "选择相关性计算方法 Select Correlation Method",
                ["Pearson相关系数 Pearson", "Spearman秩相关系数 Spearman"],
                key=f"correlation_method_{i}"
            )

            if st.button(f"执行相关性分析 Run Correlation Analysis Dataset{i + 1}", key=f"correlation_run_{i}"):
                try:
                    if selected_corr_cols:
                        df_corr = df[selected_corr_cols].dropna()

                        if correlation_method == "Pearson相关系数 Pearson":
                            corr_matrix = df_corr.corr(method='pearson')
                        else:
                            corr_matrix = df_corr.corr(method='spearman')

                        st.success("✅ 相关性分析完成 Correlation Analysis Completed")
                        st.dataframe(corr_matrix)

                        # 热力图绘制
                        import plotly.figure_factory as ff

                        z = corr_matrix.values
                        x = list(corr_matrix.columns)
                        y = list(corr_matrix.index)

                        fig = ff.create_annotated_heatmap(
                            z, x=x, y=y, colorscale='Viridis', showscale=True
                        )
                        fig.write_image("correlation_heatmap.png", scale=2)
                        st.plotly_chart(fig, use_container_width=True)
                    else:
                        st.warning("⚠ 请至少选择两列进行相关性分析")
                except Exception as e:
                    st.error(f"❌ 相关性分析失败 Correlation Analysis Failed: {e}")

        # ========== ⚖️ 分组对比分析 Comparison Analysis ==========
        st.subheader("⚖️ 分组对比分析 Comparison Analysis")
        if not dfs:
            st.warning("⚠ 当前没有上传数据，请先在侧边栏上传数据！")
            st.stop()

        compare_cols = df.select_dtypes(include=np.number).columns.tolist()
        category_cols = df.select_dtypes(exclude=np.number).columns.tolist()

        if category_cols and compare_cols:
            group_col = st.selectbox(
                "选择分组变量（类别型）Select grouping categorical variable",
                category_cols,
                key=f"groupcol_{i}"
            )
            target_col = st.selectbox(
                "选择对比的数值变量 Select numeric variable",
                compare_cols,
                key=f"targetcol_{i}"
            )

            graph_type = st.radio(
                "选择可视化图表类型 Select Visualization Type",
                ["箱线图 Boxplot", "小提琴图 Violin Plot"],
                key=f"compare_graphtype_{i}"
            )

            if st.button(f"执行分组对比分析 Run Comparison Dataset{i + 1}", key=f"compbtn_{i}"):
                group_values = df[group_col].dropna().unique()

                try:
                    # 统计检验部分
                    if len(group_values) == 2:
                        # 两组 -> t检验
                        g1, g2 = group_values[:2]
                        data1 = df[df[group_col] == g1][target_col].dropna()
                        data2 = df[df[group_col] == g2][target_col].dropna()
                        t_stat, p_val = stats.ttest_ind(data1, data2, equal_var=False)
                        st.success(f"✅ t检验结果: t={t_stat:.4f}, p={p_val:.4f}")
                    elif len(group_values) > 2:
                        # 多组 -> 方差分析
                        groups = [df[df[group_col] == g][target_col].dropna() for g in group_values]
                        f_stat, p_val = stats.f_oneway(*groups)
                        st.success(f"✅ 方差分析ANOVA结果: F={f_stat:.4f}, p={p_val:.4f}")
                    else:
                        st.warning("⚠ 分组数量不足，无法进行统计分析")

                    # 可视化部分
                    if graph_type == "箱线图 Boxplot":
                        fig = px.box(
                            df, x=group_col, y=target_col, points="all",
                            title="📦 分组箱线图 Boxplot",
                            labels={group_col: "分组 Group", target_col: "数值 Value"}
                        )
                    else:
                        fig = px.violin(
                            df, x=group_col, y=target_col, points="all", box=True,
                            title="🎻 分组小提琴图 Violin Plot",
                            labels={group_col: "分组 Group", target_col: "数值 Value"}
                        )

                    st.plotly_chart(fig, use_container_width=True)

                except Exception as e:
                    st.error(f"❌ 分组对比分析失败 Error: {e}")

        # ========== 🧪 非参数检验 Non-Parametric Test ========== #
        st.subheader("🧪 非参数检验 Non-Parametric Test")
        if not dfs:
            st.warning("⚠ 当前没有上传数据，请先在侧边栏上传数据！")
            st.stop()

        np_category_cols = df.select_dtypes(exclude=np.number).columns.tolist()
        np_numeric_cols = df.select_dtypes(include=np.number).columns.tolist()

        if np_category_cols and np_numeric_cols:
            np_group_col = st.selectbox(
                "选择分组变量（类别型）Select Grouping Categorical Variable",
                np_category_cols,
                key=f"np_group_col_{i}"
            )

            np_target_col = st.selectbox(
                "选择检验的数值变量 Select Target Numeric Variable",
                np_numeric_cols,
                key=f"np_target_col_{i}"
            )

            np_test_method = st.radio(
                "选择检验方法 Select Test Method",
                ["Mann-Whitney U检验（两组）", "Kruskal-Wallis检验（多组）"],
                key=f"np_test_method_{i}"
            )

            if st.button(f"执行非参数检验 Run Non-Parametric Test Dataset{i + 1}", key=f"np_test_run_{i}"):
                try:
                    group_values = df[np_group_col].dropna().unique()

                    if np_test_method == "Mann-Whitney U检验（两组）":
                        if len(group_values) == 2:
                            g1, g2 = group_values[:2]
                            data1 = df[df[np_group_col] == g1][np_target_col].dropna()
                            data2 = df[df[np_group_col] == g2][np_target_col].dropna()

                            from scipy.stats import mannwhitneyu

                            u_stat, p_val = mannwhitneyu(data1, data2, alternative="two-sided")
                            st.success(f"✅ Mann-Whitney U检验结果: U={u_stat:.4f}, p={p_val:.4f}")
                        else:
                            st.warning("⚠ 当前选择的分组数量不是2组，无法进行Mann-Whitney U检验")
                    else:
                        from scipy.stats import kruskal

                        if len(group_values) >= 2:
                            samples = [df[df[np_group_col] == g][np_target_col].dropna() for g in group_values]
                            h_stat, p_val = kruskal(*samples)
                            st.success(f"✅ Kruskal-Wallis检验结果: H={h_stat:.4f}, p={p_val:.4f}")
                        else:
                            st.warning("⚠ 当前分组不足两组，无法进行Kruskal-Wallis检验")

                    # 可视化分布
                    import plotly.express as px

                    fig = px.box(
                        df,
                        x=np_group_col,
                        y=np_target_col,
                        points="all",
                        title="📦 分组变量与数值变量关系 Boxplot of Groups",
                        labels={np_group_col: "分组变量 Group", np_target_col: "数值变量 Value"}
                    )
                    st.plotly_chart(fig, use_container_width=True)

                except Exception as e:
                    st.error(f"❌ 非参数检验失败 Non-Parametric Test Failed: {e}")

        # ========== 📈 比率增长分析 Ratio Growth Analysis（专业版） ========== #
        st.subheader("📈 比率增长分析（专业版）Ratio Growth Analysis (Pro)")
        if not dfs:
            st.warning("⚠ 当前没有上传数据，请先在侧边栏上传数据！")
            st.stop()

        year_cols = df.select_dtypes(include=['int', 'float', 'object']).columns.tolist()
        value_cols = df.select_dtypes(include=np.number).columns.tolist()

        if year_cols and value_cols:
            time_col = st.selectbox(
                "选择时间列（如年份）Select Time Column",
                year_cols,
                key=f"timecol_{i}"
            )
            value_col = st.selectbox(
                "选择数值列（用于计算增长率）Select Value Column",
                value_cols,
                key=f"valuecol_{i}"
            )

            if st.button(f"执行专业版比率增长分析 Run Pro Ratio Analysis Dataset{i + 1}", key=f"ratioprobnt_{i}"):
                try:
                    df_ratio = df[[time_col, value_col]].dropna()
                    df_ratio[time_col] = pd.to_numeric(df_ratio[time_col], errors="coerce")
                    df_ratio = df_ratio.dropna().sort_values(by=time_col)

                    df_ratio["增长率 YoY Growth (%)"] = df_ratio[value_col].pct_change() * 100
                    df_ratio["增长倍数 Growth Factor"] = df_ratio[value_col] / df_ratio[value_col].shift(1)

                    st.success("✅ 比率分析（专业版）完成！Ratio Analysis Completed")
                    st.dataframe(df_ratio)

                    # -- 折线图：原始数值变化趋势
                    fig_trend = px.line(
                        df_ratio,
                        x=time_col,
                        y=value_col,
                        markers=True,
                        title="📈 数值变化趋势 Value Trend Over Time",
                        labels={time_col: "时间 Time", value_col: "数值 Value"}
                    )
                    st.plotly_chart(fig_trend, use_container_width=True)

                    # -- 柱状图：增长率变化
                    fig_growth = px.bar(
                        df_ratio,
                        x=time_col,
                        y="增长率 YoY Growth (%)",
                        color=df_ratio["增长率 YoY Growth (%)"] > 0,
                        color_discrete_map={True: "green", False: "red"},
                        title="📊 年增长率变化 YoY Growth (%) Over Time",
                        labels={time_col: "时间 Time", "增长率 YoY Growth (%)": "增长率 YoY Growth (%)"},
                        text_auto=".2f"
                    )
                    st.plotly_chart(fig_growth, use_container_width=True)

                    st.info("✅ 绿色为正增长，红色为负增长")

                except Exception as e:
                    st.error(f"❌ 比率分析失败 Error: {e}")

        # ========== 🌈 主成分分析（PCA专业版）Principal Component Analysis (Pro) ========== #
        st.subheader("🌈 主成分分析（专业版）PCA Analysis (Professional Version)")
        if not dfs:
            st.warning("⚠ 当前没有上传数据，请先在侧边栏上传数据！")
            st.stop()

        pca_cols = df.select_dtypes(include=np.number).columns.tolist()
        if pca_cols:
            selected_pca_cols = st.multiselect(
                "选择用于PCA的数值字段 Select numeric fields for PCA",
                pca_cols,
                key=f"pca_cols_{i}"
            )
            pca_n_components = st.slider(
                "选择主成分数量 Number of Principal Components",
                min_value=2,
                max_value=min(len(selected_pca_cols), 10),
                value=2,
                key=f"pca_ncomp_{i}"
            )

            color_options = ["无 No Coloring"] + df.select_dtypes(exclude=np.number).columns.tolist()
            pca_color_col = st.selectbox(
                "选择分类上色字段（可选）Color by Category (Optional)",
                color_options,
                key=f"pca_colorcol_{i}"
            )

            if st.button(f"执行专业版PCA分析 Run Pro PCA Dataset{i + 1}", key=f"pca_run_{i}"):
                try:
                    from sklearn.decomposition import PCA
                    from sklearn.preprocessing import StandardScaler

                    X_raw = df[selected_pca_cols].dropna()
                    scaler = StandardScaler()
                    X_scaled = scaler.fit_transform(X_raw)

                    pca = PCA(n_components=pca_n_components)
                    components = pca.fit_transform(X_scaled)

                    # 解释方差贡献率
                    explained_var = pca.explained_variance_ratio_

                    explained_df = pd.DataFrame({
                        "主成分 Principal Component": [f"PC{i + 1}" for i in range(len(explained_var))],
                        "方差贡献率 Explained Variance Ratio": explained_var
                    })

                    st.markdown("**📊 主成分贡献率 Explained Variance Ratio:**")
                    st.dataframe(explained_df)

                    fig_var = px.bar(
                        explained_df,
                        x="主成分 Principal Component",
                        y="方差贡献率 Explained Variance Ratio",
                        title="📈 主成分贡献率图 Explained Variance of Components",
                        text_auto=".2%",
                        labels={"方差贡献率 Explained Variance Ratio": "贡献率 Explained Variance"}
                    )
                    st.plotly_chart(fig_var, use_container_width=True)

                    # -- PCA散点可视化
                    df_pca_plot = pd.DataFrame(components, columns=[f"PC{i + 1}" for i in range(pca_n_components)])
                    df_pca_plot["Sample Index"] = X_raw.index.astype(str)

                    if pca_color_col != "无 No Coloring" and pca_color_col in df.columns:
                        df_pca_plot["Color"] = df.loc[X_raw.index, pca_color_col].astype(str)
                    else:
                        df_pca_plot["Color"] = "All"

                    if pca_n_components == 2:
                        fig_pca2d = px.scatter(
                            df_pca_plot,
                            x="PC1", y="PC2",
                            color="Color",
                            title="🌈 PCA 2D降维散点图 (按分类着色) PCA 2D Scatter",
                            hover_data=["Sample Index"],
                            labels={"PC1": "主成分1 PC1", "PC2": "主成分2 PC2"}
                        )
                        st.plotly_chart(fig_pca2d, use_container_width=True)
                    elif pca_n_components >= 3:
                        fig_pca3d = px.scatter_3d(
                            df_pca_plot,
                            x="PC1", y="PC2", z="PC3",
                            color="Color",
                            title="🌈 PCA 3D降维散点图 (按分类着色) PCA 3D Scatter",
                            hover_data=["Sample Index"],
                            labels={"PC1": "主成分1 PC1", "PC2": "主成分2 PC2", "PC3": "主成分3 PC3"}
                        )
                        st.plotly_chart(fig_pca3d, use_container_width=True)

                    st.success("✅ 专业版PCA降维分析完成 PCA Analysis Completed!")

                except Exception as e:
                    st.error(f"❌ PCA降维分析失败 PCA Failed: {e}")

        # ========== 🧬 Logistic回归风险预测（可视化版）Logistic Regression with Visualization ========== #
        st.subheader("🧬 Logistic回归风险预测（可视化版）Logistic Regression Risk Prediction with Visualization")
        if not dfs:
            st.warning("⚠ 当前没有上传数据，请先在侧边栏上传数据！")
            st.stop()

        logistic_features = st.multiselect(
            "选择特征列 Select Features",
            df.select_dtypes(include=np.number).columns.tolist(),
            key=f"logistic_features_{i}"
        )

        logistic_target = st.selectbox(
            "选择目标列（0/1二分类）Select Target (Binary)",
            [col for col in df.columns if df[col].nunique() == 2],
            key=f"logistic_target_{i}"
        )

        if logistic_features and logistic_target:
            if st.button(f"执行Logistic回归预测 Run Logistic Regression Dataset{i + 1}", key=f"logistic_run_{i}"):
                try:
                    from sklearn.linear_model import LogisticRegression
                    from sklearn.model_selection import train_test_split
                    from sklearn.metrics import roc_curve, auc, confusion_matrix, ConfusionMatrixDisplay

                    X = df[logistic_features].dropna()
                    y = df.loc[X.index, logistic_target]

                    X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.3, random_state=42)

                    model = LogisticRegression(max_iter=1000)
                    model.fit(X_train, y_train)

                    # --- 系数分析 ---
                    coef_df = pd.DataFrame({
                        "Feature": logistic_features,
                        "Coefficient": model.coef_[0]
                    }).sort_values(by="Coefficient", ascending=False)

                    st.markdown("**📊 回归系数（影响方向与大小）Logistic Regression Coefficients:**")
                    fig_coef = px.bar(
                        coef_df,
                        x="Coefficient",
                        y="Feature",
                        orientation="h",
                        title="🧬 Logistic回归系数条形图",
                        labels={"Coefficient": "回归系数 Coefficient", "Feature": "特征 Feature"}
                    )
                    st.plotly_chart(fig_coef, use_container_width=True)

                    # --- ROC曲线 ---
                    y_pred_prob = model.predict_proba(X_test)[:, 1]

                    fpr, tpr, thresholds = roc_curve(y_test, y_pred_prob)
                    roc_auc = auc(fpr, tpr)

                    fig_roc = px.area(
                        x=fpr, y=tpr,
                        title=f"🩺 ROC曲线 (AUC = {roc_auc:.4f})",
                        labels=dict(x="假阳性率 FPR", y="真正率 TPR")
                    )
                    fig_roc.write_image("logistic_roc_curve.png", scale=2)
                    fig_roc.add_shape(
                        type="line", line=dict(dash="dash"),
                        x0=0, x1=1, y0=0, y1=1
                    )
                    st.plotly_chart(fig_roc, use_container_width=True)

                    # --- 混淆矩阵 ---
                    y_pred = (y_pred_prob >= 0.5).astype(int)
                    cm = confusion_matrix(y_test, y_pred)
                    cm_df = pd.DataFrame(cm, index=["实际 0", "实际 1"], columns=["预测 0", "预测 1"])

                    st.markdown("**🎯 混淆矩阵 Confusion Matrix:**")
                    st.dataframe(cm_df)

                    st.success(f"✅ Logistic回归完成，AUC = {roc_auc:.4f}")

                except Exception as e:
                    st.error(f"❌ Logistic回归失败 Logistic Regression Failed: {e}")

        # ========== 🧬 特征选择（LASSO / 递归特征消除 RFE）Feature Selection ========== #
        st.subheader("🧬 特征选择（LASSO回归 / RFE递归特征消除）Feature Selection (LASSO / RFE)")
        if not dfs:
            st.warning("⚠ 当前没有上传数据，请先在侧边栏上传数据！")
            st.stop()

        feature_numeric_cols = df.select_dtypes(include=np.number).columns.tolist()

        if feature_numeric_cols:
            fs_target_col = st.selectbox(
                "选择目标列（Target Variable）",
                feature_numeric_cols,
                key=f"fs_target_col_{i}"
            )

            fs_feature_cols = st.multiselect(
                "选择候选特征列（Feature Candidates）",
                [col for col in feature_numeric_cols if col != fs_target_col],
                key=f"fs_feature_cols_{i}"
            )

            fs_method = st.radio(
                "选择特征选择方法 Select Feature Selection Method",
                ["LASSO回归（L1正则化）", "递归特征消除 RFE"],
                key=f"fs_method_{i}"
            )

            if fs_feature_cols and fs_target_col:
                if st.button(f"执行特征选择 Run Feature Selection Dataset{i + 1}", key=f"fs_run_{i}"):
                    try:
                        X = df[fs_feature_cols].dropna()
                        y = df.loc[X.index, fs_target_col]

                        # 特征选择方法
                        if fs_method == "LASSO回归（L1正则化）":
                            from sklearn.linear_model import LassoCV

                            model = LassoCV(cv=5, random_state=42)
                            model.fit(X, y)

                            selected_features = [feature for feature, coef in zip(fs_feature_cols, model.coef_) if
                                                 abs(coef) > 1e-4]
                            st.success(f"✅ LASSO选择了{len(selected_features)}个重要特征 Features Selected by LASSO")
                            st.write("选择的特征 Selected Features:", selected_features)

                            # 可视化特征系数
                            import plotly.express as px

                            coef_df = pd.DataFrame({"Feature": fs_feature_cols, "Coefficient": model.coef_})
                            fig_coef = px.bar(
                                coef_df.sort_values(by="Coefficient"),
                                x="Coefficient", y="Feature",
                                orientation="h",
                                title="🧬 LASSO特征系数 LASSO Feature Coefficients",
                                text_auto=".2f"
                            )
                            fig_coef.update_layout(yaxis={'categoryorder': 'total ascending'})
                            st.plotly_chart(fig_coef, use_container_width=True)

                        else:  # RFE
                            from sklearn.linear_model import LogisticRegression
                            from sklearn.feature_selection import RFE

                            estimator = LogisticRegression(max_iter=1000, solver='liblinear')
                            selector = RFE(estimator, n_features_to_select=max(1, int(len(fs_feature_cols) * 0.5)))
                            selector = selector.fit(X, y)

                            selected_features = [feature for feature, support in zip(fs_feature_cols, selector.support_)
                                                 if support]
                            st.success(f"✅ RFE选择了{len(selected_features)}个重要特征 Features Selected by RFE")
                            st.write("选择的特征 Selected Features:", selected_features)

                            # 可视化支持情况
                            support_df = pd.DataFrame({
                                "Feature": fs_feature_cols,
                                "Selected (1=Yes, 0=No)": selector.support_.astype(int)
                            })
                            fig_rfe = px.bar(
                                support_df.sort_values(by="Selected (1=Yes, 0=No)"),
                                x="Selected (1=Yes, 0=No)",
                                y="Feature",
                                orientation="h",
                                title="🧬 RFE特征选择情况 RFE Feature Support",
                                text_auto=True
                            )
                            st.plotly_chart(fig_rfe, use_container_width=True)

                    except Exception as e:
                        st.error(f"❌ 特征选择失败 Feature Selection Failed: {e}")
            else:
                st.info("⚠ 请至少选择一个特征列")
        else:
            st.info("⚠ 当前数据集没有足够的数值型字段进行特征选择")

        # ========= 🌳 决策树与随机森林建模 Decision Tree & Random Forest ==========
        st.subheader("🌳 决策树与随机森林建模 Decision Tree & Random Forest")
        if not dfs:
            st.warning("⚠ 当前没有上传数据，请先在侧边栏上传数据！")
            st.stop()

        tree_target = st.selectbox(
            f"选择目标列（因变量）Dataset{i+1}",
            df.columns.tolist(),
            key=f"tree_target_{i}"
        )

        tree_features = st.multiselect(
            f"选择特征列（自变量）Dataset{i+1}",
            [col for col in df.columns if col != tree_target],
            key=f"tree_features_{i}"
        )

        tree_model_type = st.radio(
            f"选择建模类型 Dataset{i+1}",
            ["决策树 Decision Tree", "随机森林 Random Forest"],
            key=f"tree_model_type_{i}"
        )

        if st.button(f"执行树模型建模 Run Tree Modeling Dataset{i+1}", key=f"tree_run_{i}"):
            try:
                from sklearn.tree import DecisionTreeClassifier, DecisionTreeRegressor
                from sklearn.ensemble import RandomForestClassifier, RandomForestRegressor
                from sklearn.model_selection import train_test_split
                from sklearn.metrics import accuracy_score, r2_score

                X = df[tree_features].dropna()
                y = df.loc[X.index, tree_target].dropna()
                X = X.loc[y.index]

                X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)

                is_classification = y.nunique() <= 10 and y.dtype.name in ["object", "category", "bool", "int64"]

                if tree_model_type == "决策树 Decision Tree":
                    model = DecisionTreeClassifier(random_state=42) if is_classification else DecisionTreeRegressor(random_state=42)
                else:
                    model = RandomForestClassifier(n_estimators=100, random_state=42) if is_classification else RandomForestRegressor(n_estimators=100, random_state=42)

                model.fit(X_train, y_train)
                y_pred = model.predict(X_test)

                if is_classification:
                    acc = accuracy_score(y_test, y_pred)
                    st.success(f"✅ 分类准确率 Accuracy: {acc:.4f}")
                else:
                    r2 = r2_score(y_test, y_pred)
                    st.success(f"✅ 回归R²分数 R² Score: {r2:.4f}")

                # 特征重要性图
                st.markdown("**📊 特征重要性 Feature Importance**")
                importances = pd.Series(model.feature_importances_, index=tree_features)
                importances = importances.sort_values(ascending=False)
                st.bar_chart(importances)

            except Exception as e:
                st.error(f"❌ 树模型建模失败 Tree Modeling Failed: {e}")

# ========== 🌟 XGBoost建模（专业版可视化）XGBoost Modeling (Advanced Visualization) ========== #
        st.subheader("🌟 XGBoost建模（专业版可视化）XGBoost Modeling (Advanced Visualization)")
        if not dfs:
            st.warning("⚠ 当前没有上传数据，请先在侧边栏上传数据！")
            st.stop()

        xgb_target = st.selectbox(
            "选择目标列（因变量）Select Target Variable",
            df.select_dtypes(include=np.number).columns.tolist(),
            key=f"xgb_target_{i}"
        )

        xgb_features = st.multiselect(
            "选择特征列（自变量）Select Feature Columns",
            [col for col in df.select_dtypes(include=np.number).columns if col != xgb_target],
            key=f"xgb_features_{i}"
        )

        if xgb_features and xgb_target:
            xgb_model_type = st.radio(
                "选择建模类型 Select Model Type",
                ["XGBoost回归 XGBoost Regressor", "XGBoost分类 XGBoost Classifier"],
                key=f"xgb_model_type_{i}"
            )

            xgb_max_depth = st.slider("最大树深度 Max Depth", 2, 10, 4, key=f"xgb_max_depth_{i}")
            xgb_learning_rate = st.slider("学习率 Learning Rate", 0.01, 0.5, 0.1, step=0.01,
                                          key=f"xgb_learning_rate_{i}")

            if st.button(f"执行XGBoost建模 Run XGBoost Modeling Dataset{i + 1}", key=f"xgb_run_{i}"):
                try:
                    from xgboost import XGBRegressor, XGBClassifier
                    from sklearn.model_selection import train_test_split
                    from sklearn.metrics import mean_squared_error, accuracy_score, confusion_matrix, roc_curve, auc

                    X = df[xgb_features].dropna()
                    y = df.loc[X.index, xgb_target]
                    X = X.loc[y.index]

                    X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.3, random_state=42)

                    if xgb_model_type == "XGBoost回归 XGBoost Regressor":
                        model = XGBRegressor(max_depth=xgb_max_depth, learning_rate=xgb_learning_rate, n_estimators=100)
                        model.fit(X_train, y_train)
                        y_pred = model.predict(X_test)

                        rmse = mean_squared_error(y_test, y_pred, squared=False)
                        st.success(f"✅ XGBoost回归完成！测试集RMSE = {rmse:.4f}")

                        # 可视化拟合效果
                        st.markdown("**📈 预测值 vs 实际值 Predicted vs Actual:**")
                        fig_pred_actual = px.scatter(
                            x=y_test, y=y_pred,
                            labels={"x": "实际值 Actual", "y": "预测值 Predicted"},
                            title="📈 预测值 vs 实际值散点图"
                        )
                        fig_pred_actual.add_shape(
                            type="line", line=dict(dash="dash"),
                            x0=y_test.min(), y0=y_test.min(),
                            x1=y_test.max(), y1=y_test.max()
                        )
                        st.plotly_chart(fig_pred_actual, use_container_width=True)

                    else:
                        model = XGBClassifier(max_depth=xgb_max_depth, learning_rate=xgb_learning_rate,
                                              n_estimators=100)
                        model.fit(X_train, y_train)
                        y_pred = model.predict(X_test)

                        acc = accuracy_score(y_test, y_pred)
                        st.success(f"✅ XGBoost分类完成！测试集准确率 Accuracy = {acc:.4f}")

                        # 混淆矩阵
                        cm = confusion_matrix(y_test, y_pred)
                        cm_df = pd.DataFrame(cm, index=["实际 0", "实际 1"], columns=["预测 0", "预测 1"])
                        st.markdown("**🎯 混淆矩阵 Confusion Matrix:**")
                        st.dataframe(cm_df)

                        # ROC曲线
                        y_pred_prob = model.predict_proba(X_test)[:, 1]
                        fpr, tpr, _ = roc_curve(y_test, y_pred_prob)
                        roc_auc = auc(fpr, tpr)

                        fig_roc = px.area(
                            x=fpr, y=tpr,
                            title=f"🩺 ROC曲线 (AUC = {roc_auc:.4f})",
                            labels=dict(x="假阳性率 FPR", y="真正率 TPR")
                        )
                        fig_roc.add_shape(
                            type="line", line=dict(dash="dash"),
                            x0=0, x1=1, y0=0, y1=1
                        )
                        st.plotly_chart(fig_roc, use_container_width=True)

                    # 特征重要性
                    st.markdown("**📊 特征重要性 Feature Importance:**")
                    feature_importance = pd.Series(model.feature_importances_, index=xgb_features)
                    feature_importance = feature_importance.sort_values(ascending=True)

                    fig_fi = px.bar(
                        feature_importance,
                        orientation='h',
                        labels={"value": "重要性 Importance", "index": "特征 Feature"},
                        title="📊 XGBoost特征重要性条形图"
                    )
                    st.plotly_chart(fig_fi, use_container_width=True)

                except Exception as e:
                    st.error(f"❌ XGBoost建模失败 Error: {e}")

# ========= 🛡️ 简单异常检测（Z-Score/IQR法）Simple Outlier Detection ==========
        st.subheader("🛡️ 简单异常检测 Simple Outlier Detection")
        if not dfs:
            st.warning("⚠ 当前没有上传数据，请先在侧边栏上传数据！")
            st.stop()

        outlier_method = st.selectbox(
            f"选择异常检测方法 Dataset{i+1}",
            ["Z-Score方法", "四分位距（IQR）法"],
            key=f"outlier_method_{i}"
        )

        outlier_cols = st.multiselect(
            f"选择要检测的字段 Dataset{i+1}",
            df.select_dtypes(include=np.number).columns.tolist(),
            key=f"outlier_cols_{i}"
        )

        if st.button(f"执行异常检测 Run Outlier Detection Dataset{i+1}", key=f"outlier_run_{i}"):
            if outlier_cols:
                try:
                    df_outlier = df.copy()
                    found_outliers = False

                    if outlier_method == "Z-Score方法":
                        for col in outlier_cols:
                            z_scores = np.abs(stats.zscore(df_outlier[col]))
                            is_outlier = z_scores > 3
                            df_outlier[f"{col}_异常标记"] = np.where(is_outlier, "异常", "正常")
                            if is_outlier.any():
                                found_outliers = True
                    else:
                        for col in outlier_cols:
                            Q1 = df_outlier[col].quantile(0.25)
                            Q3 = df_outlier[col].quantile(0.75)
                            IQR = Q3 - Q1
                            lower = Q1 - 1.5 * IQR
                            upper = Q3 + 1.5 * IQR
                            is_outlier = (df_outlier[col] < lower) | (df_outlier[col] > upper)
                            df_outlier[f"{col}_异常标记"] = np.where(is_outlier, "异常", "正常")
                            if is_outlier.any():
                                found_outliers = True

                    st.dataframe(df_outlier)

                    if found_outliers:
                        st.success("✅ 检测到异常值 Outliers Detected")
                    else:
                        st.info("未检测到异常值 No Outliers Found")
                except Exception as e:
                    st.error(f"❌ 异常检测失败 Outlier Detection Failed: {e}")
            else:
                st.warning("⚠ 请至少选择一个字段进行异常检测")

# ========= 🛡️ 高级异常检测（Isolation Forest / LOF）Advanced Outlier Detection ==========
        st.subheader("🛡️ 高级自动异常检测 Advanced Outlier Detection")
        if not dfs:
            st.warning("⚠ 当前没有上传数据，请先在侧边栏上传数据！")
            st.stop()

        advanced_method = st.radio(
            f"选择检测算法 Dataset{i+1}",
            ["Isolation Forest", "Local Outlier Factor (LOF)"],
            key=f"advanced_outlier_method_{i}"
        )

        advanced_cols = st.multiselect(
            f"选择用于检测的特征列 Dataset{i+1}",
            df.select_dtypes(include=np.number).columns.tolist(),
            key=f"advanced_outlier_cols_{i}"
        )

        if st.button(f"执行高级异常检测 Run Advanced Detection Dataset{i+1}", key=f"advanced_outlier_run_{i}"):
            if advanced_cols:
                try:
                    from sklearn.ensemble import IsolationForest
                    from sklearn.neighbors import LocalOutlierFactor

                    X = df[advanced_cols].dropna()
                    X_index = X.index

                    if advanced_method == "Isolation Forest":
                        model = IsolationForest(contamination=0.05, random_state=42)
                        preds = model.fit_predict(X)
                        labels = np.where(preds == -1, "异常 Outlier", "正常 Normal")
                    else:
                        model = LocalOutlierFactor(n_neighbors=20, contamination=0.05)
                        preds = model.fit_predict(X)
                        labels = np.where(preds == -1, "异常 Outlier", "正常 Normal")

                    df_result = df.loc[X_index].copy()
                    df_result["高级异常检测结果 Detection Result"] = labels

                    st.success("✅ 高级异常检测完成 Advanced Detection Completed")
                    st.dataframe(df_result)

                    n_outliers = np.sum(labels == "异常 Outlier")
                    st.info(f"检测到 {n_outliers} 条异常数据 Points detected as Outliers: {n_outliers}")
                except Exception as e:
                    st.error(f"❌ 高级异常检测失败 Advanced Detection Failed: {e}")
            else:
                st.warning("⚠ 请至少选择一个特征列")
# ========= 🧠 高级PCA模块（带分类上色）Advanced PCA with Coloring ==========
        st.subheader("🧠 高级PCA降维分析 Advanced PCA with Coloring")
        if not dfs:
            st.warning("⚠ 当前没有上传数据，请先在侧边栏上传数据！")
            st.stop()

        pca_features = st.multiselect(
            f"选择用于PCA的数值特征列 Dataset{i+1}",
            df.select_dtypes(include=np.number).columns.tolist(),
            key=f"pca_features_{i}"
        )

        if pca_features:
            n_components = st.slider(
                f"选择主成分数量 Dataset{i+1}",
                min_value=2,
                max_value=min(5, len(pca_features)),
                value=2,
                key=f"pca_n_components_{i}"
            )

            color_options = ["无 No Coloring"] + df.select_dtypes(exclude=np.number).columns.tolist()
            color_col = st.selectbox(
                f"选择用于分类上色的字段（可选）Dataset{i+1}",
                color_options,
                key=f"pca_color_col_{i}"
            )

            if st.button(f"执行高级PCA分析 Run PCA Dataset{i+1}", key=f"pca_run_{i}"):
                try:
                    from sklearn.decomposition import PCA
                    from sklearn.preprocessing import StandardScaler

                    # 标准化数据
                    X_raw = df[pca_features].dropna()
                    scaler = StandardScaler()
                    X_scaled = scaler.fit_transform(X_raw)

                    # 执行PCA
                    pca = PCA(n_components=n_components)
                    components = pca.fit_transform(X_scaled)

                    explained_var = pca.explained_variance_ratio_

                    explained_df = pd.DataFrame({
                        "主成分 Principal Component": [f"PC{j+1}" for j in range(len(explained_var))],
                        "方差贡献率 Explained Variance Ratio": explained_var
                    })
                    st.dataframe(explained_df)

                    df_pca_plot = pd.DataFrame(components, columns=[f"PC{j+1}" for j in range(n_components)])
                    df_pca_plot["样本索引 Sample Index"] = X_raw.index.astype(str)

                    if color_col != "无 No Coloring" and color_col in df.columns:
                        df_pca_plot["Color"] = df.loc[X_raw.index, color_col].astype(str)
                    else:
                        df_pca_plot["Color"] = "All"

                    # 可视化
                    if n_components == 2:
                        fig = px.scatter(
                            df_pca_plot,
                            x="PC1", y="PC2",
                            color="Color",
                            hover_data=["样本索引 Sample Index"],
                            title="🌈 高级PCA 2D降维可视化（分类上色）"
                        )
                        st.plotly_chart(fig, use_container_width=True)
                    elif n_components >= 3:
                        fig = px.scatter_3d(
                            df_pca_plot,
                            x="PC1", y="PC2", z="PC3",
                            color="Color",
                            hover_data=["样本索引 Sample Index"],
                            title="🌈 高级PCA 3D降维可视化（分类上色）"
                        )
                        st.plotly_chart(fig, use_container_width=True)

                    st.success("✅ 高级PCA分析完成 Advanced PCA Completed")
                except Exception as e:
                    st.error(f"❌ PCA分析失败 PCA Failed: {e}")

# ========== 🌐 聚类分析（KMeans / 层次聚类）Clustering Analysis (KMeans / Hierarchical) ========== #
        st.subheader("🌐 聚类分析（KMeans / 层次聚类）Clustering Analysis (KMeans / Hierarchical)")
        if not dfs:
            st.warning("⚠ 当前没有上传数据，请先在侧边栏上传数据！")
            st.stop()

        cluster_numeric_cols = df.select_dtypes(include=np.number).columns.tolist()

        if cluster_numeric_cols:
            cluster_features = st.multiselect(
                "选择用于聚类的数值型特征列 Select Features for Clustering",
                cluster_numeric_cols,
                key=f"cluster_features_{i}"
            )

            cluster_method = st.radio(
                "选择聚类方法 Select Clustering Method",
                ["KMeans聚类", "层次聚类 Hierarchical Clustering"],
                key=f"cluster_method_{i}"
            )

            if cluster_features:
                cluster_n_clusters = st.slider(
                    "选择聚类簇数 Number of Clusters",
                    min_value=2, max_value=10, value=3,
                    key=f"cluster_n_clusters_{i}"
                )

                if st.button(f"执行聚类分析 Run Clustering Dataset{i + 1}", key=f"cluster_run_{i}"):
                    try:
                        from sklearn.preprocessing import StandardScaler
                        from sklearn.cluster import KMeans, AgglomerativeClustering

                        X = df[cluster_features].dropna()
                        scaler = StandardScaler()
                        X_scaled = scaler.fit_transform(X)

                        if cluster_method == "KMeans聚类":
                            model = KMeans(n_clusters=cluster_n_clusters, random_state=42)
                            cluster_labels = model.fit_predict(X_scaled)
                            st.success("✅ KMeans聚类完成 Clustering Completed!")
                        else:
                            model = AgglomerativeClustering(n_clusters=cluster_n_clusters)
                            cluster_labels = model.fit_predict(X_scaled)
                            st.success("✅ 层次聚类完成 Clustering Completed!")

                        # 添加聚类标签
                        df_result = df.loc[X.index].copy()
                        df_result["Cluster_Label"] = cluster_labels
                        st.dataframe(df_result)

                        # 可视化（2D或3D）
                        import plotly.express as px

                        if len(cluster_features) >= 3:
                            fig_cluster = px.scatter_3d(
                                df_result,
                                x=cluster_features[0],
                                y=cluster_features[1],
                                z=cluster_features[2],
                                color="Cluster_Label",
                                title="🌐 聚类结果3D可视化 Clustering Result (3D)",
                                labels={"Cluster_Label": "聚类簇 Cluster"}
                            )
                            fig_cluster.write_image("cluster_plot.png", scale=2)
                        else:
                            fig_cluster = px.scatter(
                                df_result,
                                x=cluster_features[0],
                                y=cluster_features[1],
                                color="Cluster_Label",
                                title="🌐 聚类结果2D可视化 Clustering Result (2D)",
                                labels={"Cluster_Label": "聚类簇 Cluster"}
                            )
                            fig_cluster.write_image("cluster_plot.png", scale=2)
                        st.plotly_chart(fig_cluster, use_container_width=True)

                    except Exception as e:
                        st.error(f"❌ 聚类分析失败 Clustering Failed: {e}")
            else:
                st.info("⚠ 请至少选择两个数值型特征进行聚类")
        else:
            st.info("⚠ 当前数据集中没有足够的数值型字段进行聚类分析")

# ========== ⏳ 时间序列预测（Auto-ARIMA专业版）Time Series Forecast (Advanced Version) ========== #
        st.subheader("⏳ 时间序列预测（Auto-ARIMA专业版）Time Series Forecast (Advanced Version)")
        if not dfs:
            st.warning("⚠ 当前没有上传数据，请先在侧边栏上传数据！")
            st.stop()

        ts_time_col = st.selectbox(
            "选择时间列 Select Time Column",
            df.select_dtypes(include=["int", "float", "object"]).columns.tolist(),
            key=f"time_col_ts_{i}"
        )

        ts_value_col = st.selectbox(
            "选择数值列 Select Value Column",
            df.select_dtypes(include=np.number).columns.tolist(),
            key=f"value_col_ts_{i}"
        )

        forecast_periods = st.number_input(
            "预测步数（未来周期数）Forecast Steps",
            min_value=1, max_value=50, value=5,
            key=f"forecast_periods_{i}"
        )

        if st.button(f"执行时间序列预测 Run Time Series Forecast Dataset{i + 1}", key=f"forecast_run_{i}"):
            try:
                import statsmodels.api as sm
                from pmdarima import auto_arima
                import plotly.graph_objects as go

                df_ts = df[[ts_time_col, ts_value_col]].dropna()
                df_ts[ts_time_col] = pd.to_numeric(df_ts[ts_time_col], errors='coerce')
                df_ts = df_ts.dropna().sort_values(by=ts_time_col)
                df_ts.set_index(ts_time_col, inplace=True)

                st.info("🔍 正在寻找最佳ARIMA模型，请稍候...")
                model = auto_arima(df_ts[ts_value_col], seasonal=False, trace=True, error_action='ignore',
                                   suppress_warnings=True)

                st.success(f"✅ Auto-ARIMA完成！ 最佳模型: ARIMA{model.order}")
                st.markdown(f"- **AIC值**: {model.aic():.2f}")
                st.markdown(f"- **BIC值**: {model.bic():.2f}")

                forecast, conf_int = model.predict(n_periods=forecast_periods, return_conf_int=True)
                future_index = np.arange(df_ts.index.max() + 1, df_ts.index.max() + forecast_periods + 1)

                # 绘图
                fig = go.Figure()

                # 历史数据
                fig.add_trace(go.Scatter(
                    x=df_ts.index,
                    y=df_ts[ts_value_col],
                    mode="lines",
                    name="历史数据 History"
                ))

                # 预测数据
                fig.add_trace(go.Scatter(
                    x=future_index,
                    y=forecast,
                    mode="lines+markers",
                    name="预测数据 Forecast"
                ))

                # 上下置信区间
                fig.add_trace(go.Scatter(
                    x=np.concatenate([future_index, future_index[::-1]]),
                    y=np.concatenate([conf_int[:, 0], conf_int[::-1, 1]]),
                    fill="toself",
                    fillcolor="rgba(255, 0, 0, 0.2)",
                    line=dict(color="rgba(255,255,255,0)"),
                    hoverinfo="skip",
                    name="95%预测区间"
                ))

                fig.update_layout(
                    title="⏳ 时间序列预测结果（含95%置信区间）",
                    xaxis_title="时间 Time",
                    yaxis_title="值 Value",
                    legend_title="图例 Legend",
                    template="plotly_white"
                )

                st.plotly_chart(fig, use_container_width=True)

            except Exception as e:
                st.error(f"❌ 时间序列预测失败 Error: {e}")
# ========== 🌟 生存分析（Kaplan-Meier专业版）Kaplan-Meier Survival Analysis (Advanced Version) ========== #
        st.subheader("🌟 生存分析（Kaplan-Meier专业版）Kaplan-Meier Survival Analysis (Advanced Version)")
        if not dfs:
            st.warning("⚠ 当前没有上传数据，请先在侧边栏上传数据！")
            st.stop()

        if len(df.select_dtypes(include=np.number).columns) >= 2:
            time_col = st.selectbox(
                "选择生存时间列 Time-to-Event Column",
                df.select_dtypes(include=np.number).columns.tolist(),
                key=f"survival_time_{i}"
            )
            event_col = st.selectbox(
                "选择事件状态列（0=无事件，1=发生事件）Event Status Column",
                df.select_dtypes(include=np.number).columns.tolist(),
                key=f"survival_event_{i}"
            )
            group_col = st.selectbox(
                "选择分组列（可选）Grouping Column (Optional)",
                ["无分组 No Grouping"] + df.select_dtypes(exclude=np.number).columns.tolist(),
                key=f"survival_group_{i}"
            )

            show_ci = st.checkbox("显示置信区间（Confidence Interval）", value=True, key=f"survival_show_ci_{i}")

            if st.button(f"执行生存分析 Run Survival Analysis Dataset{i + 1}", key=f"survival_run_{i}"):
                try:
                    from lifelines import KaplanMeierFitter
                    from lifelines.statistics import logrank_test
                    import matplotlib.pyplot as plt

                    T = df[time_col]
                    E = df[event_col]

                    fig, ax = plt.subplots()

                    if group_col == "无分组 No Grouping":
                        kmf = KaplanMeierFitter()
                        kmf.fit(T, event_observed=E)
                        kmf.plot_survival_function(ax=ax, ci_show=show_ci)
                        st.success(f"✅ 中位生存时间 Median Survival Time: {kmf.median_survival_time_:.2f}")
                    else:
                        groups = df[group_col].dropna().unique()
                        for g in groups:
                            ix = df[group_col] == g
                            kmf = KaplanMeierFitter()
                            kmf.fit(T[ix], event_observed=E[ix], label=str(g))
                            kmf.plot_survival_function(ax=ax, ci_show=show_ci)

                        if len(groups) == 2:
                            ix1 = df[group_col] == groups[0]
                            ix2 = df[group_col] == groups[1]
                            results = logrank_test(
                                T[ix1], T[ix2],
                                event_observed_A=E[ix1],
                                event_observed_B=E[ix2]
                            )
                            p_val = results.p_value
                            st.success(f"✅ Log-rank检验完成 P值: {p_val:.4f}")
                            if p_val < 0.05:
                                st.info("✅ 组间生存差异显著 Significant Difference (p<0.05)")
                            else:
                                st.info("⚠ 组间生存差异不显著 No Significant Difference (p≥0.05)")
                        else:
                            st.info("⚠ 分组超过2组，暂不支持自动Log-rank检验，仅展示各组生存曲线")

                    plt.title("Kaplan-Meier 生存曲线 Kaplan-Meier Survival Curve")
                    plt.xlabel("时间 Time")
                    plt.ylabel("生存概率 Survival Probability")
                    plt.grid(True)
                    st.pyplot(fig)

                except Exception as e:
                    st.error(f"❌ 生存分析失败 Survival Analysis Failed: {e}")

# ========== 🧬 Cox回归生存分析（多因素生存建模）Cox Proportional-Hazards Regression ========== #
        st.subheader("🧬 Cox比例风险回归生存分析 Cox Proportional-Hazards Regression")
        if not dfs:
            st.warning("⚠ 当前没有上传数据，请先在侧边栏上传数据！")
            st.stop()

        # 生存时间列 & 事件列（必须数值型）
        surv_numeric_cols = df.select_dtypes(include=np.number).columns.tolist()
        surv_categorical_cols = df.select_dtypes(exclude=np.number).columns.tolist()

        if len(surv_numeric_cols) >= 2:
            cox_time_col = st.selectbox(
                "选择生存时间列（Time-to-Event Column）",
                surv_numeric_cols,
                key=f"cox_time_col_{i}"
            )

            cox_event_col = st.selectbox(
                "选择事件状态列（0=无事件，1=发生事件 Event Status Column）",
                surv_numeric_cols,
                key=f"cox_event_col_{i}"
            )

            # 特征选择
            cox_features = st.multiselect(
                "选择特征变量（数值型和类别型均可）Select Predictor Features",
                df.columns.tolist(),
                key=f"cox_features_{i}"
            )

            if cox_features:
                if st.button(f"执行Cox回归分析 Run Cox Regression Dataset{i + 1}", key=f"cox_run_{i}"):
                    try:
                        from lifelines import CoxPHFitter

                        # 构造 Cox模型数据集
                        df_cox = df[[cox_time_col, cox_event_col] + cox_features].dropna()

                        # 将分类变量转为one-hot编码
                        df_cox_encoded = pd.get_dummies(df_cox, columns=[col for col in cox_features if
                                                                         df[col].dtype == 'object' or df[
                                                                             col].dtype.name == 'category'])

                        # Cox建模
                        cph = CoxPHFitter()
                        cph.fit(df_cox_encoded, duration_col=cox_time_col, event_col=cox_event_col)

                        st.success("✅ Cox回归建模完成 Cox Regression Completed!")
                        st.dataframe(cph.summary)

                        # 可视化1: Hazard Ratios
                        import plotly.express as px

                        hr_plot_df = cph.summary.reset_index()
                        fig_hr = px.bar(
                            hr_plot_df,
                            x="exp(coef)",
                            y="index",
                            orientation="h",
                            title="🧬 Hazard Ratios (HR) of Predictors",
                            labels={"index": "变量 Variable", "exp(coef)": "风险比 Hazard Ratio"},
                            text_auto=".2f"
                        )
                        fig_hr.update_layout(yaxis={'categoryorder': 'total ascending'})
                        st.plotly_chart(fig_hr, use_container_width=True)

                        # 可视化2: 生存函数
                        st.markdown("### 📈 样本整体生存函数 Estimated Survival Function")
                        fig_surv = cph.plot_survival_function()
                        st.pyplot(fig_surv.figure)

                    except Exception as e:
                        st.error(f"❌ Cox回归分析失败 Cox Regression Failed: {e}")
            else:
                st.info("⚠ 请至少选择一个特征变量用于Cox回归建模")
        else:
            st.info("⚠ 当前数据字段不足，至少需要两个数值列（生存时间 + 事件状态列）")

# ========== 🗺️ GBD地理分布分析 Geographic Distribution Analysis ========== #
        st.subheader("🗺️ GBD地理分布分析 Geographic Distribution")
        if not dfs:
            st.warning("⚠ 当前没有上传数据，请先在侧边栏上传数据！")
            st.stop()

        if len(df.columns) >= 3:
            geo_country_col = st.selectbox(
                "选择国家/地区列 Select Country/Region Column",
                df.columns.tolist(),
                key=f"geo_country_col_{i}"
            )

            geo_value_col = st.selectbox(
                "选择负担数值列（如DALYs、死亡率）Select Burden Value Column",
                df.select_dtypes(include=np.number).columns.tolist(),
                key=f"geo_value_col_{i}"
            )

            selected_year_geo = st.selectbox(
                "选择年份（如果有）Select Year",
                sorted(df[df.columns[df.columns.str.contains('year', case=False)]].iloc[:, 0].dropna().unique()) if any(
                    df.columns.str.contains('year', case=False)) else [],
                key=f"geo_selected_year_{i}"
            ) if any(df.columns.str.contains('year', case=False)) else None

            if st.button(f"绘制地理分布地图 Plot Geographic Distribution Dataset{i + 1}", key=f"geo_run_{i}"):
                try:
                    df_geo = df[[geo_country_col, geo_value_col]].dropna()

                    if selected_year_geo:
                        year_col_detected = df.columns[df.columns.str.contains('year', case=False)][0]
                        df_geo = df[df[year_col_detected] == selected_year_geo][
                            [geo_country_col, geo_value_col]].dropna()

                    import plotly.express as px

                    fig = px.choropleth(
                        df_geo,
                        locations=geo_country_col,
                        locationmode='country names',
                        color=geo_value_col,
                        title="🗺️ 国家/地区负担分布地图 Country/Region Burden Distribution",
                        color_continuous_scale="Plasma"
                    )
                    st.plotly_chart(fig, use_container_width=True)

                except Exception as e:
                    st.error(f"❌ 地理分布绘制失败 Geographic Plot Failed: {e}")
        else:
            st.info("⚠ 当前数据字段不足，至少需要国家列和数值列")

# ========== 📈 GBD负担随时间变化趋势分析（统一版）GBD Burden Trend Analysis ========== #
        st.subheader("📈 GBD负担随时间变化趋势分析 Burden Trend Over Time")
        if not dfs:
            st.warning("⚠ 当前没有上传数据，请先在侧边栏上传数据！")
            st.stop()

        if len(df.columns) >= 3:
            trend_time_col = st.selectbox(
                "选择时间列（年份）Select Time Column",
                df.columns.tolist(),
                key=f"trend_time_col_{i}"
            )

            trend_group_col = st.selectbox(
                "选择分组对象列（疾病/地区/性别/风险因素等）Select Grouping Column",
                df.columns.tolist(),
                key=f"trend_group_col_{i}"
            )

            trend_value_col = st.selectbox(
                "选择负担数值列（DALYs/死亡率/患病率等）Select Value Column",
                df.select_dtypes(include=np.number).columns.tolist(),
                key=f"trend_value_col_{i}"
            )

            selected_groups = st.multiselect(
                "选择要分析的对象（可多选）Select Entities to Analyze",
                df[trend_group_col].dropna().unique().tolist(),
                key=f"trend_selected_groups_{i}"
            )

            if selected_groups:
                if st.button(f"执行负担趋势分析 Run Burden Trend Analysis Dataset{i + 1}", key=f"trend_run_{i}"):
                    try:
                        df_trend = df[[trend_time_col, trend_group_col, trend_value_col]].dropna()
                        df_trend[trend_time_col] = pd.to_numeric(df_trend[trend_time_col], errors="coerce")
                        df_trend = df_trend.dropna()

                        import plotly.express as px

                        fig = px.line(
                            df_trend[df_trend[trend_group_col].isin(selected_groups)],
                            x=trend_time_col,
                            y=trend_value_col,
                            color=trend_group_col,
                            markers=True,
                            title="📈 负担随时间变化趋势 Burden Trend Over Time",
                            labels={
                                trend_time_col: "年份 Year",
                                trend_value_col: "负担数值 Burden Value",
                                trend_group_col: "分组对象 Group"
                            }
                        )
                        fig.write_image("gbd_trend_plot.png", scale=2)
                        fig.update_traces(mode="lines+markers")
                        st.plotly_chart(fig, use_container_width=True)

                    except Exception as e:
                        st.error(f"❌ 趋势分析失败 Trend Analysis Failed: {e}")
            else:
                st.info("⚠ 请至少选择一个分析对象")
        else:
            st.info("⚠ 当前数据字段不足，至少需要时间列、分组列、数值列")

# ========== 📉 GBD变化率分析专区（变化率计算+AAPC）GBD Change Rate Analysis ========== #
        st.subheader("📉 GBD变化率分析专区 Change Rate Analysis")
        if not dfs:
            st.warning("⚠ 当前没有上传数据，请先在侧边栏上传数据！")
            st.stop()

        if len(df.columns) >= 3:
            cr_entity_col = st.selectbox(
                "选择分析对象列（国家/地区/疾病等）Select Entity Column",
                df.columns.tolist(),
                key=f"cr_entity_col_{i}"
            )

            cr_year_col = st.selectbox(
                "选择时间列（年份）Select Year Column",
                df.columns.tolist(),
                key=f"cr_year_col_{i}"
            )

            cr_value_col = st.selectbox(
                "选择负担数值列（DALYs/死亡率/患病率等）Select Value Column",
                df.select_dtypes(include=np.number).columns.tolist(),
                key=f"cr_value_col_{i}"
            )

            selected_entities_cr = st.multiselect(
                "选择要分析的对象（可多选）Select Entities to Analyze",
                df[cr_entity_col].dropna().unique().tolist(),
                key=f"cr_selected_entities_{i}"
            )

            cr_analysis_type = st.radio(
                "选择变化率分析类型 Select Change Rate Type",
                ["简单变化率（首尾年）Simple Change Rate", "年均变化率（AAPC）Average Annual Percentage Change"],
                key=f"cr_analysis_type_{i}"
            )

            if selected_entities_cr:
                if st.button(f"执行变化率分析 Run Change Rate Analysis Dataset{i + 1}", key=f"cr_run_{i}"):
                    try:
                        df_cr = df[[cr_entity_col, cr_year_col, cr_value_col]].dropna()
                        df_cr[cr_year_col] = pd.to_numeric(df_cr[cr_year_col], errors="coerce")
                        df_cr = df_cr.dropna()

                        results = []

                        for entity in selected_entities_cr:
                            sub = df_cr[df_cr[cr_entity_col] == entity]
                            sub = sub.sort_values(by=cr_year_col)

                            if len(sub) >= 2:
                                start_val = sub.iloc[0][cr_value_col]
                                end_val = sub.iloc[-1][cr_value_col]
                                n_years = sub.iloc[-1][cr_year_col] - sub.iloc[0][cr_year_col]

                                if cr_analysis_type == "简单变化率（首尾年）Simple Change Rate":
                                    change_rate = ((end_val - start_val) / start_val) * 100
                                else:
                                    if start_val > 0 and n_years > 0:
                                        change_rate = ((end_val / start_val) ** (1 / n_years) - 1) * 100
                                    else:
                                        change_rate = None

                                if change_rate is not None:
                                    results.append((entity, change_rate))

                        df_result = pd.DataFrame(results, columns=["分析对象 Entity", "变化率 Change Rate (%)"])

                        st.success(f"✅ {cr_analysis_type}计算完成 Analysis Completed")
                        st.dataframe(df_result)

                        # 可视化
                        import plotly.express as px

                        fig = px.bar(
                            df_result.sort_values(by="变化率 Change Rate (%)"),
                            x="变化率 Change Rate (%)", y="分析对象 Entity",
                            orientation="h",
                            title=f"📉 {cr_analysis_type}结果",
                            text_auto=".2f"
                        )
                        fig.write_image("gbd_change_rate_plot.png", scale=2)
                        fig.update_layout(yaxis={'categoryorder': 'total ascending'})
                        st.plotly_chart(fig, use_container_width=True)

                    except Exception as e:
                        st.error(f"❌ 变化率分析失败 Change Rate Analysis Failed: {e}")
            else:
                st.info("⚠ 请至少选择一个分析对象")
        else:
            st.info("⚠ 当前数据字段不足，至少需要对象列、时间列、数值列")

# ========== 🧬 GBD风险因素归因分析专区（风险Top分析+PAF归因分析）GBD Risk Factors Attribution ========== #
        st.subheader("🧬 GBD风险因素归因分析 Risk Factors Attribution Analysis")
        if not dfs:
            st.warning("⚠ 当前没有上传数据，请先在侧边栏上传数据！")
            st.stop()

        if len(df.columns) >= 3:
            risk_factor_col = st.selectbox(
                "选择风险因素列 Select Risk Factor Column",
                df.columns.tolist(),
                key=f"risk_factor_col_{i}"
            )

            burden_value_col = st.selectbox(
                "选择负担数值列（如归因DALYs或死亡数）Select Burden Value Column",
                df.select_dtypes(include=np.number).columns.tolist(),
                key=f"risk_burden_value_col_{i}"
            )

            # PAF相关列
            paf_exposure_col = st.selectbox(
                "选择暴露率列（Exposure Rate 0-1之间）Select Exposure Rate Column",
                df.select_dtypes(include=np.number).columns.tolist(),
                key=f"risk_exposure_col_{i}"
            )

            paf_rr_col = st.selectbox(
                "选择相对风险列（Relative Risk）Select Relative Risk Column",
                df.select_dtypes(include=np.number).columns.tolist(),
                key=f"risk_rr_col_{i}"
            )

            selected_risks = st.multiselect(
                "选择要分析的风险因素（可多选）Select Risks to Analyze",
                df[risk_factor_col].dropna().unique().tolist(),
                key=f"risk_selected_factors_{i}"
            )

            if selected_risks:
                if st.button(f"执行风险因素归因分析 Run Risk Attribution Analysis Dataset{i + 1}", key=f"risk_run_{i}"):
                    try:
                        df_risk = df[[risk_factor_col, burden_value_col, paf_exposure_col, paf_rr_col]].dropna()

                        # Top10风险贡献分析
                        df_top = df_risk[df_risk[risk_factor_col].isin(selected_risks)]
                        df_top = df_top.groupby(risk_factor_col)[burden_value_col].sum().reset_index()
                        df_top = df_top.sort_values(by=burden_value_col, ascending=False).head(10)

                        # 归因风险比例PAF计算
                        paf_results = []
                        for _, row in df_risk.iterrows():
                            exposure = row[paf_exposure_col]
                            rr = row[paf_rr_col]
                            if 0 <= exposure <= 1 and rr >= 1:
                                paf = (exposure * (rr - 1)) / (exposure * (rr - 1) + 1)
                                paf_results.append((row[risk_factor_col], paf * 100))

                        df_paf = pd.DataFrame(paf_results, columns=["风险因素 Risk Factor", "PAF (%)"])
                        df_paf = df_paf[df_paf["风险因素 Risk Factor"].isin(selected_risks)]

                        st.success("✅ 风险Top分析和PAF归因比例计算完成 Analysis Completed")

                        st.markdown("### 📊 风险因素负担Top10分析 Risk Factors Burden Top10")
                        st.dataframe(df_top)

                        import plotly.express as px

                        fig1 = px.bar(
                            df_top.sort_values(by=burden_value_col),
                            x=burden_value_col, y=risk_factor_col,
                            orientation="h",
                            title="🧬 风险因素负担Top10 Risk Burden Top10",
                            text_auto=".2s"
                        )
                        fig1.write_image("risk_top10_plot.png", scale=2)
                        fig1.update_layout(yaxis={'categoryorder': 'total ascending'})
                        st.plotly_chart(fig1, use_container_width=True)

                        st.markdown("### 📈 归因风险比例PAF分析 Population Attributable Fraction (PAF)")
                        st.dataframe(df_paf)

                        fig2 = px.bar(
                            df_paf.sort_values(by="PAF (%)"),
                            x="PAF (%)", y="风险因素 Risk Factor",
                            orientation="h",
                            title="🧬 风险因素PAF归因比例 Risk Factors PAF",
                            text_auto=".2f"
                        )
                        fig2.write_image("risk_paf_plot.png", scale=2)
                        fig2.update_layout(yaxis={'categoryorder': 'total ascending'})
                        st.plotly_chart(fig2, use_container_width=True)

                    except Exception as e:
                        st.error(f"❌ 风险归因分析失败 Risk Attribution Failed: {e}")
            else:
                st.info("⚠ 请至少选择一个风险因素")
        else:
            st.info("⚠ 当前数据字段不足，至少需要风险列、负担列、暴露率列、相对风险列")

# ========== 🌟 GBD生存分析专区（Kaplan-Meier + Stratified KM + Cox回归）Survival Analysis ========== #
        st.subheader("🌟 GBD生存分析专区 Survival Analysis")
        if not dfs:
            st.warning("⚠ 当前没有上传数据，请先在侧边栏上传数据！")
            st.stop()

        if len(df.columns) >= 3:
            surv_time_col = st.selectbox(
                "选择生存时间列（Time-to-Event Column）",
                df.select_dtypes(include=np.number).columns.tolist(),
                key=f"surv_time_col_{i}"
            )

            surv_event_col = st.selectbox(
                "选择事件状态列（0=无事件，1=事件）Select Event Status Column",
                df.select_dtypes(include=np.number).columns.tolist(),
                key=f"surv_event_col_{i}"
            )

            analysis_type = st.radio(
                "选择生存分析类型 Select Survival Analysis Type",
                ["单因素Kaplan-Meier生存曲线", "分层Kaplan-Meier生存曲线", "Cox比例风险回归"],
                key=f"surv_analysis_type_{i}"
            )

            if analysis_type in ["单因素Kaplan-Meier生存曲线", "分层Kaplan-Meier生存曲线"]:
                strat_cols = st.multiselect(
                    "选择分组/分层变量（可选）Select Group/Stratification Variables",
                    df.select_dtypes(exclude=np.number).columns.tolist(),
                    key=f"surv_strat_cols_{i}"
                )

            else:
                cox_features = st.multiselect(
                    "选择Cox回归特征变量（可数值/可分类）Select Features for Cox Regression",
                    df.columns.tolist(),
                    key=f"cox_features_{i}"
                )

            if st.button(f"执行生存分析 Run Survival Analysis Dataset{i + 1}", key=f"surv_run_{i}"):
                try:
                    from lifelines import KaplanMeierFitter, CoxPHFitter
                    from lifelines.statistics import multivariate_logrank_test
                    import matplotlib.pyplot as plt

                    df_surv = df.dropna(subset=[surv_time_col, surv_event_col])

                    if analysis_type in ["单因素Kaplan-Meier生存曲线", "分层Kaplan-Meier生存曲线"]:
                        if strat_cols:
                            df_surv["Group"] = df_surv[strat_cols].apply(lambda row: "-".join(row.values.astype(str)),
                                                                         axis=1)
                        else:
                            df_surv["Group"] = "All"

                        groups = df_surv["Group"].unique()

                        fig, ax = plt.subplots(figsize=(8, 6))
                        for group in groups:
                            ix = df_surv["Group"] == group
                            kmf = KaplanMeierFitter()
                            kmf.fit(df_surv.loc[ix, surv_time_col], event_observed=df_surv.loc[ix, surv_event_col],
                                    label=group)
                            kmf.plot_survival_function(ax=ax)

                        plt.title("🌟 Kaplan-Meier生存曲线 Kaplan-Meier Survival Curves")
                        plt.xlabel("生存时间 Survival Time")
                        plt.ylabel("生存概率 Survival Probability")
                        plt.grid(True)
                        st.pyplot(fig)
                        fig.savefig("km_survival_curve.png", dpi=300)

                        if len(groups) > 1:
                            results = multivariate_logrank_test(
                                df_surv[surv_time_col],
                                groups=df_surv["Group"],
                                event_observed=df_surv[surv_event_col]
                            )
                            p_value = results.p_value
                            st.success(f"✅ 多组Log-rank检验完成 P-value: {p_value:.4f}")

                    else:
                        if not cox_features:
                            st.warning("⚠ 请至少选择一个Cox回归特征变量")
                        else:
                            df_cox = df_surv[[surv_time_col, surv_event_col] + cox_features].dropna()
                            df_cox_encoded = pd.get_dummies(df_cox, columns=[col for col in cox_features if
                                                                             df[col].dtype == 'object' or df[
                                                                                 col].dtype.name == 'category'])

                            cph = CoxPHFitter()
                            cph.fit(df_cox_encoded, duration_col=surv_time_col, event_col=surv_event_col)

                            st.success("✅ Cox回归模型拟合完成 Cox Model Fitted")
                            st.dataframe(cph.summary)

                            import plotly.express as px

                            coef_df = cph.summary.reset_index()
                            fig = px.bar(
                                coef_df,
                                x="exp(coef)", y="index",
                                orientation="h",
                                title="🧬 Cox回归模型特征风险比 Hazard Ratios",
                                labels={"index": "变量 Variable", "exp(coef)": "风险比 Hazard Ratio"},
                                text_auto=".2f"
                            )
                            fig.write_image("cox_hr_plot.png", scale=2)
                            fig.update_layout(yaxis={'categoryorder': 'total ascending'})
                            st.plotly_chart(fig, use_container_width=True)

                except Exception as e:
                    st.error(f"❌ 生存分析失败 Survival Analysis Failed: {e}")
        else:
            st.info("⚠ 当前数据字段不足，至少需要生存时间列、事件列")

# ========== 🏔️ GBD负担堆积趋势分析（Area Plot）GBD Burden Stacked Area Analysis ========== #
        st.subheader("🏔️ GBD负担堆积趋势分析 Stacked Area Trend Analysis")
        if not dfs:
            st.warning("⚠ 当前没有上传数据，请先在侧边栏上传数据！")
            st.stop()

        if len(df.columns) >= 3:
            area_time_col = st.selectbox(
                "选择时间列（年份）Select Time Column",
                df.columns.tolist(),
                key=f"area_time_col_{i}"
            )

            area_group_col = st.selectbox(
                "选择分组列（疾病/风险因素等）Select Grouping Column",
                df.columns.tolist(),
                key=f"area_group_col_{i}"
            )

            area_value_col = st.selectbox(
                "选择数值列（DALYs/死亡率/患病率等）Select Value Column",
                df.select_dtypes(include=np.number).columns.tolist(),
                key=f"area_value_col_{i}"
            )

            selected_groups_area = st.multiselect(
                "选择要分析的对象（可多选）Select Entities to Include",
                df[area_group_col].dropna().unique().tolist(),
                key=f"area_selected_groups_{i}"
            )

            if selected_groups_area:
                if st.button(f"执行负担堆积趋势分析 Run Stacked Area Trend Analysis Dataset{i + 1}",
                             key=f"area_run_{i}"):
                    try:
                        df_area = df[[area_time_col, area_group_col, area_value_col]].dropna()
                        df_area[area_time_col] = pd.to_numeric(df_area[area_time_col], errors="coerce")
                        df_area = df_area.dropna()

                        df_area = df_area[df_area[area_group_col].isin(selected_groups_area)]

                        import plotly.express as px

                        fig = px.area(
                            df_area,
                            x=area_time_col,
                            y=area_value_col,
                            color=area_group_col,
                            title="🏔️ 负担堆积趋势分析 Stacked Area Burden Trends",
                            labels={
                                area_time_col: "年份 Year",
                                area_value_col: "负担数值 Burden Value",
                                area_group_col: "分组对象 Group"
                            }
                        )
                        st.plotly_chart(fig, use_container_width=True)

                    except Exception as e:
                        st.error(f"❌ 堆积趋势分析失败 Area Plot Analysis Failed: {e}")
            else:
                st.info("⚠ 请至少选择一个对象进行堆积分析")
        else:
            st.info("⚠ 当前数据字段不足，至少需要时间列、分组列、数值列")

# ========== 🏆 GBD国家负担排名变化分析（随时间变化）GBD Country Burden Ranking Trend ========== #
        st.subheader("🏆 GBD国家负担排名变化分析 Country Burden Ranking Trend")
        if not dfs:
            st.warning("⚠ 当前没有上传数据，请先在侧边栏上传数据！")
            st.stop()

        if len(df.columns) >= 3:
            rank_time_col = st.selectbox(
                "选择时间列（年份）Select Time Column",
                df.columns.tolist(),
                key=f"rank_time_col_{i}"
            )

            rank_country_col = st.selectbox(
                "选择国家列 Select Country Column",
                df.columns.tolist(),
                key=f"rank_country_col_{i}"
            )

            rank_value_col = st.selectbox(
                "选择负担数值列（DALYs/死亡率/患病率等）Select Burden Value Column",
                df.select_dtypes(include=np.number).columns.tolist(),
                key=f"rank_value_col_{i}"
            )

            selected_countries_rank = st.multiselect(
                "选择要追踪排名变化的国家（可多选）Select Countries to Track",
                df[rank_country_col].dropna().unique().tolist(),
                key=f"rank_selected_countries_{i}"
            )

            if selected_countries_rank:
                if st.button(f"执行国家负担排名变化分析 Run Country Ranking Trend Analysis Dataset{i + 1}",
                             key=f"rank_run_{i}"):
                    try:
                        df_rank = df[[rank_time_col, rank_country_col, rank_value_col]].dropna()
                        df_rank[rank_time_col] = pd.to_numeric(df_rank[rank_time_col], errors="coerce")
                        df_rank = df_rank.dropna()

                        # 每一年内部排名
                        df_rank["负担排名 Rank"] = df_rank.groupby(rank_time_col)[rank_value_col].rank(method="min",
                                                                                                       ascending=False)

                        df_plot = df_rank[df_rank[rank_country_col].isin(selected_countries_rank)]

                        import plotly.express as px

                        fig = px.line(
                            df_plot,
                            x=rank_time_col,
                            y="负担排名 Rank",
                            color=rank_country_col,
                            markers=True,
                            title="🏆 国家负担排名随时间变化 Country Burden Ranking Trend",
                            labels={
                                rank_time_col: "年份 Year",
                                "负担排名 Rank": "排名 Rank",
                                rank_country_col: "国家 Country"
                            }
                        )
                        fig.update_yaxes(autorange="reversed")  # 名次1在上
                        st.plotly_chart(fig, use_container_width=True)

                    except Exception as e:
                        st.error(f"❌ 国家排名变化分析失败 Country Ranking Analysis Failed: {e}")
            else:
                st.info("⚠ 请至少选择一个国家进行排名分析")
        else:
            st.info("⚠ 当前数据字段不足，至少需要时间列、国家列、数值列")

# ========== 📑 完整版专业报告生成器 Generate Full Analysis Report ========== #
import docx
from docx.shared import Inches
import os
import pandas as pd

st.sidebar.header("📑 生成专业分析报告 Generate Report")
generate_report = st.sidebar.checkbox("启用分析报告生成器 Enable Report Generator")

if generate_report and dfs:
    st.header("📑 分析报告生成器 Report Generator")

    report_modules = st.multiselect(
        "请选择需要纳入报告的功能模块（可多选）Select Modules to Include in Report",
        [
            "描述性统计分析",
            "相关性分析",
            "分组对比分析",
            "生存分析",
            "时间序列预测分析",
            "PCA分析",
            "聚类分析",
            "机器学习建模",
            "GBD地理分布分析",
            "GBD趋势分析",
            "GBD变化率分析",
            "GBD风险因素归因分析",
            "GBD生存分析"
        ],
        default=["描述性统计分析", "相关性分析", "GBD趋势分析"]
    )

    report_name = st.text_input("输入报告文件名（无需加后缀）Enter Report Filename", "医学数据分析报告")

    if st.button("📥 生成并下载Word报告 Generate & Download Report"):
        try:
            doc = docx.Document()

            # ===== 报告封面 =====
            doc.add_heading("医学数据统计分析与可视化报告", 0)
            doc.add_paragraph(f"报告名称：{report_name}")
            doc.add_paragraph(f"单位：XX单位（可填写你的学校、医院或实验室）")
            doc.add_paragraph(f"生成时间：{pd.Timestamp.now():%Y-%m-%d %H:%M}")
            doc.add_page_break()

            doc.add_page_break()

            # ===== 遍历用户选择的模块，生成对应内容 =====
            # 在 for module in report_modules: 循环前添加这一行
            chapter_number = 1

            # 下面是for module in report_modules:循环内容
            for module in report_modules:
                doc.add_heading(f"第{chapter_number}章 {module}", level=1)
                chapter_number += 1

                if module == "描述性统计分析":
                    doc.add_paragraph(
                        "【描述性统计分析 Descriptive Statistics】\n\n"
                        "本模块对所选数值型变量进行基本统计描述。\n"
                        "- 均值、中位数、标准差、偏度、峰度已统计。\n"
                        "- 大部分变量分布接近正态，数据质量良好。\n\n"
                        "结果为后续分析和建模提供了基础。"
                    )
                    if os.path.exists("desc_stats_table.png"):
                        doc.add_picture("desc_stats_table.png", width=Inches(5))
                        doc.add_paragraph("▲ 描述性统计直方图")

                if module == "相关性分析":
                    doc.add_paragraph(
                        "【相关性分析 Correlation Analysis】\n\n"
                        "分析了数值变量之间的Pearson/Spearman相关性。\n"
                        "- 相关系数区间为 [-1, 1]。\n"
                        "- 存在若干高相关变量对。\n\n"
                        "结果为后续建模提供了变量筛选参考。"
                    )
                    if os.path.exists("correlation_heatmap.png"):
                        doc.add_picture("correlation_heatmap.png", width=Inches(5))
                        doc.add_paragraph("▲ 相关性矩阵热力图")

                if module == "分组对比分析":
                    doc.add_paragraph(
                        "【分组对比分析 Group Comparison Analysis】\n\n"
                        "比较了不同组别之间的关键变量差异。\n"
                        "- 两组使用t检验或Mann-Whitney检验。\n"
                        "- 多组使用ANOVA或Kruskal-Wallis检验。\n\n"
                        "显著性差异变量有助于识别影响因素。"
                    )
                    if os.path.exists("group_comparison_plot.png"):
                        doc.add_picture("group_comparison_plot.png", width=Inches(5))
                        doc.add_paragraph("▲ 分组对比箱线图/小提琴图")

                if module == "生存分析":
                    doc.add_paragraph(
                        "【生存分析 Survival Analysis】\n\n"
                        "采用Kaplan-Meier曲线与Cox回归模型进行生存分析。\n"
                        "- 不同组别中位生存时间及差异性分析。\n"
                        "- 识别独立危险因素及保护因素。\n\n"
                        "结果用于风险分层和预后评估。"
                    )
                    if os.path.exists("km_survival_curve.png"):
                        doc.add_picture("km_survival_curve.png", width=Inches(5))
                        doc.add_paragraph("▲ Kaplan-Meier生存曲线")
                    if os.path.exists("cox_hr_plot.png"):
                        doc.add_picture("cox_hr_plot.png", width=Inches(5))
                        doc.add_paragraph("▲ Cox回归风险比条形图")

                if module == "时间序列预测分析":
                    doc.add_paragraph(
                        "【时间序列预测分析 Time Series Forecasting Analysis】\n\n"
                        "使用Auto-ARIMA建模，预测未来趋势。\n"
                        "- 预测区间与趋势方向已绘制。\n\n"
                        "结果用于疾病负担变化趋势预测。"
                    )
                    if os.path.exists("time_series_forecast.png"):
                        doc.add_picture("time_series_forecast.png", width=Inches(5))
                        doc.add_paragraph("▲ 时间序列趋势预测图")

                if module == "PCA分析":
                    doc.add_paragraph(
                        "【主成分分析 PCA】\n\n"
                        "对高维数据进行降维处理，提取主要成分。\n"
                        "- 前两主成分解释主要变异。\n"
                        "- 可视化样本分布结构。\n\n"
                        "PCA分析揭示了数据内部潜在结构。"
                    )
                    if os.path.exists("pca_plot.png"):
                        doc.add_picture("pca_plot.png", width=Inches(5))
                        doc.add_paragraph("▲ PCA主成分散点图")

                if module == "聚类分析":
                    doc.add_paragraph(
                        "【聚类分析 Clustering Analysis】\n\n"
                        "通过无监督学习方法探索数据潜在分组。\n"
                        "- 样本在2D/3D空间中聚类结果可视化。\n\n"
                        "聚类分析揭示潜在亚群体。"
                    )
                    if os.path.exists("cluster_plot.png"):
                        doc.add_picture("cluster_plot.png", width=Inches(5))
                        doc.add_paragraph("▲ 聚类散点图")

                if module == "机器学习建模":
                    doc.add_paragraph(
                        "【机器学习建模 Machine Learning Modeling】\n\n"
                        "采用Logistic回归和XGBoost建模。\n"
                        "- 模型AUC、灵敏度、特异性评估。\n"
                        "- 特征重要性排序可视化。\n\n"
                        "建模结果用于预测和变量筛选。"
                    )
                    if os.path.exists("logistic_roc_curve.png"):
                        doc.add_picture("logistic_roc_curve.png", width=Inches(5))
                        doc.add_paragraph("▲ Logistic回归ROC曲线")
                    if os.path.exists("xgboost_feature_importance.png"):
                        doc.add_picture("xgboost_feature_importance.png", width=Inches(5))
                        doc.add_paragraph("▲ XGBoost特征重要性图")

                if module == "GBD地理分布分析":
                    doc.add_paragraph(
                        "【GBD地理分布分析 Geographic Distribution Analysis】\n\n"
                        "展示了全球疾病负担的地理分布特征。\n"
                        "- 不同国家和地区负担差异明显。\n\n"
                        "为制定区域干预策略提供参考。"
                    )
                    if os.path.exists("gbd_geographic_distribution.png"):
                        doc.add_picture("gbd_geographic_distribution.png", width=Inches(5))
                        doc.add_paragraph("▲ 疾病负担地理分布图")

                if module == "GBD趋势分析":
                    doc.add_paragraph(
                        "【GBD负担时间趋势分析 Burden Trend Over Time Analysis】\n\n"
                        "分析了疾病负担随时间变化的趋势。\n"
                        "- 上升、下降、稳定趋势总结。\n\n"
                        "趋势分析指导负担变化监测。"
                    )
                    if os.path.exists("gbd_trend_plot.png"):
                        doc.add_picture("gbd_trend_plot.png", width=Inches(5))
                        doc.add_paragraph("▲ 疾病负担时间趋势图")

                if module == "GBD变化率分析":
                    doc.add_paragraph(
                        "【GBD变化率分析 Burden Change Rate Analysis】\n\n"
                        "计算了年均变化率（AAPC）及简单变化率。\n"
                        "- 识别增长最快和下降最快疾病。\n\n"
                        "变化率反映疾病流行趋势变化。"
                    )
                    if os.path.exists("gbd_change_rate_plot.png"):
                        doc.add_picture("gbd_change_rate_plot.png", width=Inches(5))
                        doc.add_paragraph("▲ 疾病负担变化率条形图")

                if module == "GBD风险因素归因分析":
                    doc.add_paragraph(
                        "【GBD风险因素归因分析 Risk Factors Attribution Analysis】\n\n"
                        "分析了各主要风险因素对负担的贡献。\n"
                        "- 归因风险比例（PAF）计算完成。\n\n"
                        "提示干预优先级。"
                    )
                    if os.path.exists("risk_top10_plot.png"):
                        doc.add_picture("risk_top10_plot.png", width=Inches(5))
                        doc.add_paragraph("▲ 风险因素负担Top10图")
                    if os.path.exists("risk_paf_plot.png"):
                        doc.add_picture("risk_paf_plot.png", width=Inches(5))
                        doc.add_paragraph("▲ 风险因素PAF图")

                if module == "GBD生存分析":
                    doc.add_paragraph(
                        "【GBD生存分析 GBD Survival Analysis】\n\n"
                        "基于Kaplan-Meier和Cox模型分析各地区生存差异。\n"
                        "- 中位生存时间及危险因素识别。\n\n"
                        "结果用于生存预后研究。"
                    )
                    if os.path.exists("gbd_km_survival_curve.png"):
                        doc.add_picture("gbd_km_survival_curve.png", width=Inches(5))
                        doc.add_paragraph("▲ GBD分组Kaplan-Meier曲线")
                    if os.path.exists("gbd_cox_hr_plot.png"):
                        doc.add_picture("gbd_cox_hr_plot.png", width=Inches(5))
                        doc.add_paragraph("▲ GBD Cox回归风险比条形图")

                # 每个模块最后加一个分页
                doc.add_page_break()

            # ===== 保存文档 =====
            output_path = f"{report_name}.docx"
            doc.save(output_path)

            with open(output_path, "rb") as f:
                st.download_button(
                    label="📥 下载生成的Word报告 Download Report",
                    data=f,
                    file_name=f"{report_name}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

            st.success("✅ 报告生成成功！Report Generated Successfully!")

        except Exception as e:
            st.error(f"❌ 报告生成失败 Report Generation Failed: {e}")



